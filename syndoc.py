import os
import sys
import time
import watchdog.events
import watchdog.observers

from uuid import uuid4
from urllib.parse import quote

import collections.abc
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from progressbar import AnimatedMarker, Bar, BouncingBar, Counter, ETA, \
    AdaptiveETA, FileTransferSpeed, FormatLabel, Percentage, \
    ProgressBar, ReverseBar, RotatingMarker, \
    SimpleProgress, Timer, UnknownLength


def dir2json(path):
    base = os.path.basename(path)
    d = {
        'id': str(uuid4()),
        'name': base
    }
    if os.path.isdir(path):
        d['type'] = "directory"
        d['child'] = [dir2json(os.path.join(path, x))
                      for x in os.listdir(path)]
    else:
        pdf = f'https://raw.githubusercontent.com/yilmazchef/powershell-notebooks/main/Notebooks/English/{quote(base)}'
        md = f'https://raw.githubusercontent.com/yilmazchef/powershell-notebooks/main/Notebooks/English/{quote(base.replace(".md", ".pdf"))}'

        d['type'] = "file"
        d['link'] = md,
        d['download'] = pdf

    return d


def only4md(folder_path: str, file_type: str) -> list:
    paths = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(file_type.lower()):
                paths.append(os.path.join(root, file))

    return paths


def img2base64(image_file, output_file):

    # need base 64
    import base64
    import sys
    # open the image
    image = open(image_file, 'rb')
    # read it
    image_read = image.read()
    # encode it as base 64
    # after python>=3.9, use `encodebytes` instead of `encodestring`
    image_64_encode = base64.encodestring(image_read) if sys.version_info < (
        3, 9) else base64.encodebytes(image_read)
    # convert the image base 64 to a string
    image_string = str(image_64_encode)
    # replace the newline characters
    image_string = image_string.replace("\\n", "")
    # replace the initial binary
    image_string = image_string.replace("b'", "")
    # replace the final question mark
    image_string = image_string.replace("'", "")
    # add the image tags
    image_string = '<p><img src="data:image/png;base64,' + image_string + '"></p>'
    # write it out
    image_result = open(output_file, 'w')
    image_result.write(image_string)


def ipynb2md(ipynb_file: str) -> str:
    cmdlet = f"jupyter nbconvert --to markdown \"{ipynb_file}\""
    os.system(cmdlet)
    return ipynb_file.replace(".ipynb", ".md")


def md2ipynb(md_file: str) -> str:

    ipynb_file = md_file.replace(".md", ".ipynb")
    cmdlet = f"pandoc \"{md_file}\" -o \"{ipynb_file}\""
    os.system(cmdlet)
    return ipynb_file


def md2odt(md_file):
    odt_file = md_file.replace(".md", "odt")
    cmdlet = f"pandoc -t odt \"{md_file}\" -o \"{odt_file}\""
    os.system(cmdlet)


def md2pptx(md_file: str) -> str:

    c = collections
    c.abc = collections.abc

    pptx_path = md_file.replace(".md", ".pptx")
    cmdlet = f"pandoc -V fontsize=12pt \"{md_file}\" -s --wrap auto -o \"{pptx_path}\""
    os.system(cmdlet)

    return pptx_path


def md2docx(md_file: str) -> str:

    docx_file = md_file.replace(".md", ".docx")
    cmdlet = f"pandoc \"{md_file}\" -f markdown -o \"{docx_file}\""
    os.system(cmdlet)
    return docx_file


def h4docx(docx_file, header_image, header_text=None):

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        header = document.sections[0].header

        htable = header.add_table(1, 2, Inches(4))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture(header_image, width=Inches(3))

        if header_text is not None:
            ht1 = htab_cells[1].add_paragraph(header_text)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # saving the blank document
        document.save(docx_file)


def f4docx(docx_file, footer_image, footer_text=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        footer = document.sections[0].footer

        ftable = footer.add_table(1, 2, Inches(4))
        ftab_cells = ftable.rows[0].cells
        ft0 = ftab_cells[0].add_paragraph()
        fh = ft0.add_run()

        imgSize = Inches(4)

        if footer_text is not None:
            imgSize = Inches(3)
            ft1 = ftab_cells[1].add_paragraph(footer_text)
            ft1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fh.add_picture(footer_image, width=imgSize)

        document.save(docx_file)


def md2pdf(md_file: str) -> str:

    pdf_file = md_file.replace(".md", ".pdf")
    cmdlet = f"pandoc \"{md_file}\" --pdf-engine=xelatex -o \"{pdf_file}\""
    os.system(cmdlet)
    return pdf_file


def update_all():

    src_path = input(
        "Source folder: "
    )

    if src_path == "":
        src = sys.argv[1] if len(sys.argv) > 1 else os.getcwd()

    source_file_list = only4md(src_path, ".md")

    pBarMax = len(source_file_list)
    widgets = [Percentage(),
               ' ', Bar(),
               ' ', ETA(),
               ' ', AdaptiveETA()]
    pBar = ProgressBar(widgets=widgets, maxval=pBarMax)

    pBar.start()
    pBarCount = 0
    for source_file in source_file_list:

        docx_file = md2docx(source_file)
        h4docx(docx_file, os.path.join(os.getcwd(), "Templates", "header.png"), open(
            os.path.join(os.getcwd(), "templates", "header.txt"), 'r'))
        f4docx(docx_file, os.path.join(os.getcwd(), "Templates", "footer.png"), open(
            os.path.join(os.getcwd(), "templates", "footer.txt"), 'r'))

        pptx_file = md2pptx(source_file)

        pdf_file = md2pdf(source_file)

        pBarCount += 1
        pBar.update(pBarCount)

    pBar.finish()


class Handler(watchdog.events.PatternMatchingEventHandler):
    def __init__(self):
        # Set the patterns for PatternMatchingEventHandler
        watchdog.events.PatternMatchingEventHandler.__init__(self, patterns=['*.md'],
                                                             ignore_directories=True, case_sensitive=False)
        print(sys.getdefaultencoding())

    def on_created(self, event):
        print(f"{event.src_path} is created.")
        # Event is created, you can process it now

    def on_modified(self, event):
        docx_file = md2docx(event.src_path)
        h4docx(docx_file, os.path.join(os.getcwd(), "Templates", "header.png"), open(
            os.path.join(os.getcwd(), "templates", "header.txt"), 'r'))
        f4docx(docx_file, os.path.join(os.getcwd(), "Templates", "footer.png"), open(
            os.path.join(os.getcwd(), "templates", "footer.txt"), 'r'))
        pptx_file = md2pptx(event.src_path)
        pdf_file = md2pdf(event.src_path)

    def on_deleted(self, event):
        os.remove(event.src_path.replace(".md", "docx"))
        os.remove(event.src_path.replace(".md", "pptx"))
        os.remove(event.src_path.replace(".md", "pdf"))


if __name__ == "__main__":

    src_path = sys.argv[1] if len(sys.argv) > 1 else os.getcwd()
    event_handler = Handler()
    observer = watchdog.observers.Observer()
    observer.schedule(event_handler, path=src_path, recursive=True)
    observer.start()
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()
